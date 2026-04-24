from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 创建文档
doc = Document()

# 添加标题
title = doc.add_heading('数据开发规范文档', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加副标题
subtitle = doc.add_paragraph('文档生成时间：2026-03-27')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].italic = True

doc.add_paragraph()  # 空行

# 一、命名规范
doc.add_heading('一、命名规范', level=1)
doc.add_paragraph('1. 表命名：ods_源系统_表名_desc', style='List Number')
doc.add_paragraph('2. 字段命名：使用小写字母，下划线分隔', style='List Number')
doc.add_paragraph('3. 任务命名：job_业务域_功能描述', style='List Number')

# 二、开发流程
doc.add_heading('二、开发流程', level=1)
doc.add_paragraph('1. 需求分析', style='List Number')
doc.add_paragraph('2. 设计评审', style='List Number')
doc.add_paragraph('3. 开发实现', style='List Number')
doc.add_paragraph('4. 测试验证', style='List Number')
doc.add_paragraph('5. 上线发布', style='List Number')

# 三、代码规范
doc.add_heading('三、代码规范', level=1)
doc.add_paragraph('1. SQL 语句格式化', style='List Number')
doc.add_paragraph('2. 添加必要注释', style='List Number')
doc.add_paragraph('3. 避免硬编码', style='List Number')

# 四、质量管控
doc.add_heading('四、质量管控', level=1)
doc.add_paragraph('1. 数据校验规则', style='List Number')
doc.add_paragraph('2. 异常监控告警', style='List Number')
doc.add_paragraph('3. 质量报告生成', style='List Number')

# 五、权限管理
doc.add_heading('五、权限管理', level=1)
doc.add_paragraph('1. 最小权限原则', style='List Number')
doc.add_paragraph('2. 定期权限审查', style='List Number')
doc.add_paragraph('3. 操作日志审计', style='List Number')

# 六、文档要求
doc.add_heading('六、文档要求', level=1)
doc.add_paragraph('1. 设计文档完整', style='List Number')
doc.add_paragraph('2. 接口文档清晰', style='List Number')
doc.add_paragraph('3. 运维手册齐全', style='List Number')

# 添加页脚说明
doc.add_paragraph()
doc.add_paragraph('来源：图片 OCR 识别', style='Intense Quote')

# 保存文档
doc.save('/home/rskuser/.openclaw/workspace/D:\\Users\\yr.ju.CN1\\.openclaw\\workspace/数据开发规范文档.docx')
print('Word 文档创建成功！')
